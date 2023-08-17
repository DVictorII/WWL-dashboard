import matplotlib.pyplot as plt
import matplotlib.dates as mdates
import matplotlib.lines as mlines
import uuid


import threading
import time

import math

from flask import Blueprint, jsonify, request, session
from flask_cors import cross_origin
from datetime import datetime, timedelta, date
from sqlalchemy import text, create_engine, inspect
import numpy as np

from app import db, dict_helper, app

from docx import Document
from docx.shared import Cm

import routes.wwl_functions as wwl


from routes.piezometers_data import piezometer_details
import os
import psycopg2


dbname = "wwlengineering_rossing"
user = "WWL_ADMIN"
password = "WWL#2023"
host = "wwl-rossing.crnkanilun4m.ap-southeast-2.rds.amazonaws.com"
dev = True

report_status = "off"
report_filename_global = ""


def dbconnect():
    conn = psycopg2.connect(host=host, database=dbname, user=user, password=password)
    return conn


tables = [
    "node_64031_1",
    "node_64218_1",
    "node_64047_1",
    "node_64038_1",
    "node_64038_2",
    "node_64050_1",
    "node_64148_1",
    "node_64037_1",
    "node_64037_2",
    "node_64068_1",
    "node_64053_1",
    "node_64033_1",
    "node_64033_2",
    "position",
    "task_details",
    "task_list",
    "users",
    "incident_reports",
    "last_readings",
    "node_10230_1",
    "node_10262_1",
    "node_10262_2",
    "node_10262_3",
    "node_10262_4",
    "node_10317_1",
    "node_10317_2",
    "node_16301_1",
    "node_16329_1",
    "node_16444_1",
    "node_16546_1",
    "node_16564_1",
    "node_16567_1",
    "node_16596_1",
    "node_16601_1",
    "node_16629_1",
    "node_16725_1",
    "node_16752_1",
    "node_16761_1",
    "node_1815_1",
    "node_1815_2",
    "node_1815_3",
    "node_1815_4",
    "node_1815_5",
    "node_1816_1",
    "node_1816_2",
    "node_1816_3",
    "node_1816_4",
    "node_1816_5",
    "node_1817_1",
    "node_1817_2",
    "node_1817_3",
    "node_1817_4",
    "node_1817_5",
    "node_1819_1",
    "node_1819_2",
    "node_1819_3",
    "node_1819_4",
    "node_1819_5",
    "node_1821_1",
    "node_1821_2",
    "node_1821_3",
    "node_1821_4",
    "node_1821_5",
    "node_1849_1",
    "node_1849_2",
    "node_1849_3",
    "node_1849_4",
    "node_1849_5",
    "node_1853_1",
    "node_1853_2",
    "node_1853_3",
    "node_1853_4",
    "node_1853_5",
    "node_1876_1",
    "node_1876_2",
    "node_1876_3",
    "node_1876_4",
    "node_1876_5",
    "node_20618_1",
    "node_20696_1",
    "node_20764_1",
    "node_20775_1",
    "node_2087_1",
    "node_2087_2",
    "node_2087_3",
    "node_2087_4",
    "node_2087_5",
    "node_20892_1",
    "node_20989_1",
    "node_21008_1",
    "node_21024_1",
    "node_2115_1",
    "node_2115_2",
    "node_2121_1",
    "node_2121_2",
    "node_2137_1",
    "node_2137_2",
    "node_2137_3",
    "node_2137_4",
    "node_2137_5",
    "node_2423_1",
    "node_2423_2",
    "node_2423_3",
    "node_2423_4",
    "node_2423_5",
    "node_3006_1",
    "node_3023_1",
    "node_3031_1",
    "node_3060_1",
    "node_9548_1",
    "node_9548_2",
    "node_9852_1",
    "node_9926_1",
    "piezometer_data",
    "piezometer_details",
    "piezometer_reports",
]


def build_word_report(app1, piezos, reqDate):
    # print("PIEZOS", piezos)
    # print("REQDATE", reqDate)
    with app1.app_context():
        document = Document()

        sections = document.sections

        for section in sections:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(1.91)
            section.right_margin = Cm(1.91)

        document.add_heading("Appendix: Piezometer Information")
        document.add_paragraph("")
        document.add_picture(
            os.path.abspath(
                f"../client/public/sectionReport/sections/sectionOverview.jpg"
            ),
            width=Cm(18),
            height=Cm(12),
        )

        document.add_page_break()

        paddockList = list(set(list(map(lambda x: x["paddock"], piezos))))

        finalArr = []

        # CREATE SECTION CHART IMAGES
        plot_section_chart()

        for i in range(0, len(paddockList)):
            paddock = paddockList[i]

            # PUT PADDOCK TITLE INTO PDF, START NEW SECTION
            p = document.add_paragraph("")
            p.add_run(f"{i+1}. Paddock: {paddock}").bold = True

            ################################################################

            sectionsList = list(
                filter(
                    lambda z: z != None and z != "?",
                    list(
                        set(
                            list(
                                map(
                                    lambda x: x["section"],
                                    list(
                                        filter(
                                            lambda y: y["paddock"] == paddock, piezos
                                        )
                                    ),
                                )
                            )
                        )
                    ),
                )
            )

            for j in range(0, len(sectionsList)):
                section = sectionsList[j]

                # PENDING ADD SECTION MAP AND START NEW PART IN THE PDF HERE

                ps = document.add_paragraph("")
                ps.add_run(f"{i+1}.{j+1} {section.replace('-', ' ', 1)}").bold = True

                # AÑADIR LOS MAPAS PARA TODAS LAS SECCIONES EXCEPTO LA 7-8A Y 9A (TODAVÍA NO LAS TENEMOS, QUEDA PENDIENTE AÑADIRLAS)
                if (
                    section.lower() != "section-7-8a"
                    and section.lower() != "section-9a"
                ):
                    document.add_picture(
                        os.path.abspath(
                            f"../client/public/sectionReport/maps/{section.lower()}.png"
                        ),
                        width=Cm(15.24),
                        height=Cm(9.4),
                    )
                    document.add_paragraph(" ")

                ################################################################

                piezo_list_by_section = list(
                    filter(lambda y: y["section"] == section, piezos)
                )

                for k in range(0, len(piezo_list_by_section)):
                    piezometer = piezo_list_by_section[k]

                    if k == 0:
                        filtered_list = list(
                            filter(
                                lambda x: x["section"] == section
                                and x["datalogger"] != None
                                and x["channel"] != None,
                                piezos,
                            )
                        )

                        first_piezo_with_datalogger_and_channel = filtered_list[0]

                        # BUILD SECTION CHART WITH THE FIRST PIEZOMETER THAT HAS DATALOGGER AND CHANNEL

                        document.add_picture(
                            os.path.abspath(
                                f"../client/public/sectionReport/sections/{section.lower() }.png"
                            ),
                            width=Cm(17),
                            height=Cm(8.5),
                        )
                        document.add_page_break()
                        #################################################################################

                    if piezometer["status"] == 1:
                        ps = document.add_paragraph(f"{piezometer['id']}")

                        if (
                            piezometer["datalogger"] != None
                            and piezometer["channel"] != None
                            and (
                                f"node_{piezometer['datalogger']}_{piezometer['channel']}"
                                in tables
                            )
                            # and (piezometer["paddock"] == "CROWN")
                        ):
                            filename = plot_readings_chart(piezometer, 90, reqDate)
                            print("FILENAME", filename)

                            document.add_picture(
                                os.path.abspath(
                                    f"../client/public/sectionReport/readings/{filename}"
                                ),
                                width=Cm(17),
                                height=Cm(8.5),
                            )

                            document.add_paragraph(" ")

                        else:
                            p_no_readings = document.add_paragraph("")
                            p_no_readings.add_run(
                                f"No readings for the current piezometer"
                            ).bold = True

                document.add_page_break()

        # GET PIEZO READINGS FOR EACH PIEZOMETER IN EACH SECTION

        global report_filename_global

        report_filename_global = f"word_report_{str(uuid.uuid4())}.docx"

        document.save(
            os.path.abspath(f"../client/dist/report_word/{report_filename_global}")
        )
        document.save(
            os.path.abspath(f"../client/public/report_word/{report_filename_global}")
        )

        print("REPORT FINISHED")

        global report_status
        report_status = "ok"


def move_duplicates(x):
    result = []
    seen = set()
    for value in x:
        if value in seen:
            new_value = value + 2
            while new_value in seen:
                new_value += 2
            seen.add(new_value)
            result.append(new_value)
        else:
            seen.add(value)
            result.append(value)
    return result


def plot_section_chart():
    natural_ground = "data/sections/natural_ground/"
    new_ground = "data/sections/new_ground/"
    test = wwl.get_data_by_section(natural_ground, new_ground)

    for i, paddock in enumerate(test, start=1):
        # img = Image(paddock['image'], width=200, height=150)

        for j, section in enumerate(paddock["sections"], start=1):
            plt.figure(figsize=(16, 6))
            plt.tight_layout()
            x = [item[0] for item in section["coordinates"]]
            y1 = [item[1] for item in section["coordinates"]]
            y2 = [item[2] for item in section["coordinates"]]
            nodes = sorted(section["piezometers"], key=lambda x: x[2])
            piezometer_x = [item[2] for item in nodes]
            piezometer_x = move_duplicates(piezometer_x)
            piezometer_location = [item[3] for item in nodes]
            piezometer_reading = [item[4] for item in nodes]
            piezometer_name = [item[0] for item in nodes]
            piezometer_status = [item[1] for item in nodes]
            # Add items to legend
            blue_circle = mlines.Line2D(
                [], [], color="blue", marker="o", linestyle="None", label="Active"
            )
            blue_ncircle = mlines.Line2D(
                [],
                [],
                color="blue",
                marker="o",
                linestyle="None",
                label="Active with negative reading",
                fillstyle="none",
            )
            red_circle = mlines.Line2D(
                [], [], color="red", marker="o", linestyle="None", label="Damaged"
            )
            yellow_circle = mlines.Line2D(
                [],
                [],
                color="yellow",
                marker="o",
                linestyle="None",
                label="Disconnected",
            )
            lgd = plt.legend(
                handles=[blue_circle, blue_ncircle, red_circle, yellow_circle], loc=3
            )
            plt.plot(x, y1, color="#7b8831", label="Original Ground (RLm)")
            plt.plot(x, y2, color="#876538", label="Tailing Surface (RLm)")
            index = 1
            for i, reading in enumerate(piezometer_reading):
                if piezometer_status[i] == 1:
                    color = "blue"
                elif piezometer_status[i] == 2:
                    color = "red"
                elif piezometer_status[i] == 3:
                    color = "yellow"

                if piezometer_status[i] == 1:
                    if piezometer_reading[i] < piezometer_location[i]:
                        tag = (
                            str(index) + ") " + piezometer_name[i] + " - Negative"
                        )  # + str(reading)
                        plt.scatter(
                            piezometer_x[i],
                            piezometer_location[i] - 1,
                            marker="v",
                            color=color,
                            facecolors="none",
                        )
                        plt.plot(
                            [piezometer_x[i], piezometer_x[i]],
                            [piezometer_location[i], piezometer_location[i] - 1],
                            color=color,
                            linestyle="dashed",
                        )
                        plt.scatter(
                            piezometer_x[i],
                            piezometer_location[i],
                            marker="o",
                            color=color,
                            label=tag,
                            facecolors="none",
                        )
                    else:
                        tag = (
                            str(index)
                            + ") "
                            + piezometer_name[i]
                            + " - "
                            + str(reading)
                            + " RLm"
                        )
                        plt.scatter(
                            piezometer_x[i],
                            reading,
                            marker="v",
                            color=color,
                            facecolors="none",
                        )
                        plt.plot(
                            [piezometer_x[i], piezometer_x[i]],
                            [piezometer_location[i], reading],
                            color=color,
                        )
                        plt.scatter(
                            piezometer_x[i],
                            piezometer_location[i],
                            marker="o",
                            color=color,
                            label=tag,
                        )
                    plt.annotate(
                        "(" + str(index) + ")",
                        (piezometer_x[i], piezometer_location[i]),
                        textcoords="offset points",
                        xytext=(5, 5),
                        ha="center",
                        fontsize=8,
                    )
                    index += 1
                elif piezometer_status[i] == 2 or piezometer_status[i] == 3:
                    tag = str(index) + ") " + piezometer_name[i]
                    plt.scatter(
                        piezometer_x[i],
                        piezometer_location[i],
                        marker="o",
                        color=color,
                        label=tag,
                    )
                    plt.annotate(
                        "(" + str(index) + ")",
                        (piezometer_x[i], piezometer_location[i]),
                        textcoords="offset points",
                        xytext=(5, 5),
                        ha="center",
                        fontsize=8,
                    )
                    index += 1

            plt.xlabel("Chainage (m)")
            plt.ylabel("Elevation (m)")
            plt.title(paddock["name"] + " - " + section["name"], fontsize=18)
            plt.legend()
            plt.gca().add_artist(lgd)
            plt.xticks(range(min(x), max(x) + 1, 50))
            plt.yticks(range(int(min(y1)), int(max(y2) + 1), 5))
            plt.grid(True)
            plt.margins(x=0)

            plt.savefig(
                f"../client/public/sectionReport/sections/{section['name'].lower()}.png",
                bbox_inches="tight",
            )
            # elements.append(Paragraph(section['name'], getSampleStyleSheet()['Heading2'], keepTogether=True))
            plt.close()


# def plot_section_chart(piezometer):
#     natural_ground = "data/sections/natural_ground/"
#     new_ground = "data/sections/new_ground/"

#     nodes, data, section = wwl.get_data_by_section(
#         piezometer["datalogger"], piezometer["channel"], natural_ground, new_ground
#     )

#     nodes = sorted(nodes, key=lambda x: x[2])
#     x = [item[0] for item in data]
#     y1 = [item[1] for item in data]
#     y2 = [item[2] for item in data]

#     piezometer_x = [item[2] for item in nodes]
#     piezometer_location = [item[3] for item in nodes]
#     piezometer_reading = [item[4] for item in nodes]
#     piezometer_name = [item[0] for item in nodes]
#     piezometer_status = [item[1] for item in nodes]

#     plt.figure(figsize=(16, 6))

#     plt.plot(x, y1, color="#7b8831", label="Original ground")
#     plt.plot(x, y2, color="#876538", label="Tailing surface")

#     # plt.scatter(
#     #     piezometer_x, piezometer_location, color="red", label="Piezometer location"
#     # )
#     # plt.scatter(
#     #     piezometer_x, piezometer_reading, color="blue", label="Current PWP level"
#     # )

#     index = 1
#     for i, reading in enumerate(piezometer_reading):
#         print("READING", reading)
#         if piezometer_status[i] == 1:
#             color = "blue"
#         elif piezometer_status[i] == 2:
#             color = "red"
#         elif piezometer_status[i] == 3:
#             color = "yellow"

#         if piezometer_status[i] == 1:
#             tag = str(index) + ") " + piezometer_name[i] + " - " + str(reading) + " KPa"
#             plt.scatter(piezometer_x[i], reading, marker="v", color=color)
#             plt.scatter(
#                 piezometer_x[i],
#                 piezometer_location[i],
#                 marker="o",
#                 color=color,
#                 label=tag,
#             )
#             if int(reading) == 0:
#                 plt.plot(
#                     [piezometer_x[i], piezometer_x[i]],
#                     color=color,
#                 )
#             else:
#                 plt.plot(
#                     [piezometer_x[i], piezometer_x[i]],
#                     [piezometer_location[i], reading],
#                     color=color,
#                 )
#             # plt.text(piezometer_x[i], piezometer_location[i], "("+ str(index) + ")" , ha='center', va='bottom',bbox=dict(facecolor='white', alpha=0.5))
#             plt.annotate(
#                 "(" + str(index) + ")",
#                 (piezometer_x[i], piezometer_location[i]),
#                 textcoords="offset points",
#                 xytext=(5, 5),
#                 ha="center",
#                 fontsize=8,
#             )
#             index += 1
#         elif piezometer_status[i] == 2 or piezometer_status[i] == 3:
#             tag = str(index) + ") " + piezometer_name[i]
#             plt.scatter(
#                 piezometer_x[i],
#                 piezometer_location[i],
#                 marker="o",
#                 color=color,
#                 label=tag,
#             )
#             plt.annotate(
#                 "(" + str(index) + ")",
#                 (piezometer_x[i], piezometer_location[i]),
#                 textcoords="offset points",
#                 xytext=(5, 5),
#                 ha="center",
#                 fontsize=8,
#             )
#             index += 1

#     plt.xlabel("Chainage (m)")
#     plt.ylabel("Elevation (m)")
#     plt.title(f"Graph of {piezometer['section'].replace('-',' ',1)} ")

#     plt.legend()
#     plt.xticks(range(min(x), max(x) + 1, 50))
#     plt.yticks(range(int(min(y1)), int(max(y2) + 1), 5))
#     plt.grid(True)

#     # Save chart picture
#     file_path_to_save = os.path.abspath(
#         f"../client/public/sectionReport/sections/{piezometer['section'].lower()}.png"
#     )

#     plt.savefig(file_path_to_save, bbox_inches="tight")

#     plt.close("all")

#     filename = f"{piezometer['section'].lower()}.png"

#     return filename


def plot_readings_chart(piezometer, daysAgo, reqDate):
    print("PIEZO", piezometer, daysAgo, reqDate)
    arr = reqDate.split("-")
    intArr = list(map(lambda x: int(x), arr))

    recentDate = date(*intArr).strftime("%Y-%m-%d 00:00:00")
    d = date(*intArr) - timedelta(days=int(90))
    pastDate = d.strftime("%Y-%m-%d 00:00:00")

    result = db.session.execute(
        text(
            f"SELECT time,pressure FROM node_{piezometer['datalogger']}_{piezometer['channel']} WHERE time >= '{pastDate}' AND time <= '{recentDate}' ;"
        )
    )

    lectures = [dict(r._mapping) for r in result]

    if len(lectures) < 500:
        number = math.ceil(len(lectures) / 500)

    else:
        if len(lectures) > 8000:
            number = len(lectures) // 800

        elif len(lectures) > 4000:
            number = len(lectures) // 650

        else:
            number = len(lectures) // 500

    def get_only_some_lectures(idx_and_item):
        index, item = idx_and_item

        if index % number == 0:
            return item
        else:
            return "none"

    timeArr = list(
        map(
            lambda x: x["time"].strftime("%Y-%m-%d %H:%M:%S"),
            list(
                filter(
                    lambda y: y != "none",
                    list(map(lambda x: get_only_some_lectures(x), enumerate(lectures))),
                )
            ),
        )
    )

    pressureArr = list(
        map(
            lambda x: x["pressure"],
            list(
                filter(
                    lambda y: y != "none",
                    list(map(lambda x: get_only_some_lectures(x), enumerate(lectures))),
                )
            ),
        )
    )

    t = timeArr
    s = pressureArr

    if len(pressureArr) != 0:
        maxLimit = max(pressureArr)

        minLimit = min(pressureArr)
    else:
        maxLimit = 0
        minLimit = 100

    def testFunc(idx_and_item):
        index, item = idx_and_item

        span = round(len(t) / 25)

        if index % span == 0:
            return item
        else:
            return ""

    spacedTime = list(map(testFunc, enumerate(t)))

    fig = plt.figure(figsize=(16, 6))

    ax = fig.add_subplot(1, 1, 1)
    ax.tick_params(length=0)

    plt.ylim([minLimit - 10, maxLimit + 40])

    plt.plot(t, s, label="Pressure readings (KPa)")

    plt.legend()
    plt.xticks(np.arange(len(spacedTime)), spacedTime, rotation=45)
    plt.xlabel("Dates")
    plt.ylabel("Pressure (KPa)")

    arrPastDate = pastDate.split(" ")
    arrRecentDate = recentDate.split(" ")

    plt.title(
        f"{piezometer['id']} - {piezometer['section']} - {piezometer['paddock']} - {arrPastDate[0]} to {arrRecentDate[0]} "
    )
    plt.gca().yaxis.grid(True)

    if len(pressureArr) != 0:
        plt.fill_between(t, s, minLimit - 10, color=["#477C9A"], alpha=0.1)

    file_path = os.path.abspath(
        f"../client/public/sectionReport/readings/{piezometer['datalogger']}_{piezometer['channel']}.png"
    )

    filename = f"{piezometer['datalogger']}_{piezometer['channel']}.png"

    print("PREPARING FILENAME", filename)
    plt.savefig(file_path, bbox_inches="tight")

    plt.close()

    return filename


# except:
#     return "no readings"


section_chart_routes = Blueprint("section_chart_routes", __name__)


@section_chart_routes.route("/api/v1/paddock-chart", methods=["POST"])
@cross_origin()
def build_paddocks_information_chart():
    global report_status
    report_status = "pending"

    result = piezometer_details.query.all()

    print("got piezos")

    piezos = dict_helper(result)

    print("processed piezos, entering build word function")

    reqDate = request.json["date"]

    threading.Thread(target=build_word_report, args=(app, piezos, reqDate)).start()

    # build_word_report(piezos, reqDate)
    return jsonify({"message": "Creating report", "status": "pending"})


@section_chart_routes.route("/api/v1/report-status", methods=["GET"])
@cross_origin()
def get_report_status():
    global report_status
    return jsonify({"status": report_status})


@section_chart_routes.route("/api/v1/download-word-report", methods=["GET"])
@cross_origin()
def download_word_report():
    global report_filename_global

    return jsonify({"filename": report_filename_global})


@section_chart_routes.route("/api/v1/sections-info", methods=["GET"])
@cross_origin()
def get_wuwu():
    natural_ground = "data/sections/natural_ground/"
    new_ground = "data/sections/new_ground/"

    paddock_list = wwl.get_data_by_section(natural_ground, new_ground)

    return jsonify(
        {
            "list": paddock_list,
        }
    )
